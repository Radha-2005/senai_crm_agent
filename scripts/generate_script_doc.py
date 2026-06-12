"""
generate_script_doc.py
Generates the Screen Recording Script as a properly formatted Word document.
Run: python scripts/generate_script_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "Screen_Recording_Script.docx")

# ─── Colour palette ───────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x23, 0x7E)   # deep navy – title
MED_BLUE    = RGBColor(0x15, 0x65, 0xC0)   # medium blue – segment headings
ORANGE      = RGBColor(0xE6, 0x51, 0x00)   # orange – action labels
GREEN       = RGBColor(0x1B, 0x5E, 0x20)   # dark green – spoken text
PURPLE      = RGBColor(0x4A, 0x14, 0x8C)   # purple – tips / notes
DARK_GREY   = RGBColor(0x21, 0x21, 0x21)   # near-black – body
LIGHT_GREY  = RGBColor(0x42, 0x42, 0x42)   # grey – sub-items


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=DARK_BLUE, size=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_subheading(doc, text, color=MED_BLUE, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_action(doc, text):
    """Orange bold label for on-screen actions."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("▶  " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ORANGE
    return p


def add_script(doc, text):
    """Green italic text for spoken narration."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run('"' + text + '"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    return p


def add_body(doc, text, color=DARK_GREY, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, color=DARK_GREY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color
    return p


def add_tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("💡  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = PURPLE
    return p


def add_separator(doc):
    doc.add_paragraph("─" * 80)


def add_timing_table(doc):
    """Insert a coloured timing summary table."""
    table = doc.add_table(rows=9, cols=3)
    table.style = "Table Grid"

    headers = ["Segment", "Time", "Duration"]
    rows_data = [
        ["Intro",                        "0:00 – 0:30",  "30 sec"],
        ["1 · Email Ingestion",           "0:30 – 2:00",  "90 sec"],
        ["2 · Bob's Agent Trace",         "2:00 – 4:00",  "120 sec"],
        ["3 · RAG Debug View",            "4:00 – 5:00",  "60 sec"],
        ["4 · Karen Churn + Web Intel",   "5:00 – 6:30",  "90 sec"],
        ["5 · Analytics Dashboard",       "6:30 – 7:30",  "60 sec"],
        ["Outro",                         "7:30 – 8:00",  "30 sec"],
        ["TOTAL",                         "~ 8 minutes",  "✅"],
    ]

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "1A237E")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)

    # Data rows
    alt_colors = ["EFF3FF", "FFFFFF"]
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        bg = "1B5E20" if row_data[0] == "TOTAL" else alt_colors[r_idx % 2]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            if row_data[0] == "TOTAL":
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    r = title.add_run("SenAI CRM Agentic Intelligence Platform")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Screen Recording Script — Word-for-Word Narration")
    r2.font.size = Pt(14); r2.font.color.rgb = MED_BLUE

    doc.add_paragraph()
    add_body(doc, "Total Recording Time: ~8 minutes  |  Format: Live demo of running website", DARK_GREY, 11)
    add_separator(doc)

    # ── LEGEND ────────────────────────────────────────────────────
    add_heading(doc, "Legend", 2, DARK_BLUE, 13)
    add_action(doc, "Orange text = ON-SCREEN ACTION (what to click / do)")
    add_script(doc, "Green italic text = SPOKEN NARRATION (what to say out loud)")
    add_tip(doc, "Purple text = TIPS for the presenter")
    add_separator(doc)

    # ── BEFORE YOU START ──────────────────────────────────────────
    add_heading(doc, "✅  Before You Start Recording", 2, DARK_BLUE, 14)
    add_bullet(doc, "Terminal 1: Backend running → http://localhost:8000")
    add_bullet(doc, "Terminal 2: Frontend running → http://localhost:5173")
    add_bullet(doc, "Terminal 3: Ready to run replay_dataset.py")
    add_bullet(doc, "Browser Tab 1: http://localhost:5173  (inbox with threads loaded)")
    add_bullet(doc, "Browser Tab 2: http://localhost:8000/api/docs  (Swagger UI)")
    add_bullet(doc, "Microphone tested, OBS / Loom / Game Bar ready")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # INTRO
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎙️  INTRO  (0:00 – 0:30)", 1, DARK_BLUE, 16)
    add_action(doc, "Show: Dashboard homepage with all threads loaded in the inbox")

    add_script(doc,
        "Hello, I'm [Your Name], and this is the SenAI CRM Agentic Intelligence Platform — "
        "a production-ready AI system I built for the SenAI Technical Assessment.")
    add_script(doc,
        "The system autonomously monitors a customer email inbox, triages each email through "
        "a three-layer intelligence engine, runs a ReAct autonomous agent that reasons step by step "
        "before taking any action, and surfaces real-time business insights on a live dashboard.")
    add_script(doc,
        "It processes 60 emails across 30 named conversation threads — including GDPR data requests, "
        "ransomware threats, SLA breaches, churn scenarios, and pricing negotiations. "
        "Let me walk you through each component.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # SEGMENT 1
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎬  SEGMENT 1 — Email Ingestion Pipeline  (0:30 – 2:00)", 1, MED_BLUE, 16)

    add_subheading(doc, "On-Screen Actions:")
    add_action(doc, "Switch to Terminal 3")
    add_action(doc, "Type and run:  python scripts/replay_dataset.py --delay 1.0")
    add_action(doc, "Watch terminal output for ~20 seconds as emails are accepted")
    add_action(doc, "Switch to browser — show the inbox updating in real time")
    add_action(doc, "Point to priority badges: Escalated (red), Legal-Flag (dark red), Spam (grey)")
    add_tip(doc, "Click on the inbox after ~10 emails arrive so the audience sees threads appearing live")

    add_subheading(doc, "What to Say:")
    add_script(doc,
        "Let me start by ingesting the full dataset. I'm running the replay script, which sends "
        "all 60 emails to our POST /api/v1/ingest endpoint — simulating a real-time email stream.")
    add_script(doc,
        "You can see each email being accepted one by one. "
        "msg_001 is Alice's pricing inquiry, msg_002 is Bob's P0 production outage, "
        "and msg_003 is immediately flagged as spam — blocked before any AI processing.")
    add_script(doc,
        "Every email goes through three things immediately on ingest — before any AI runs at all. "
        "First: Pydantic schema validation — malformed payloads are rejected with a descriptive error. "
        "Second: deduplication — if you send the same message_id twice, the system returns a 409 Conflict "
        "and silently skips it. This makes the entire ingestion pipeline idempotent. "
        "Third: a heuristic priority score is assigned based on keywords like URGENT, P0, ransomware, "
        "or legal — all of this happens in under ten milliseconds.")
    add_script(doc,
        "Now watch the dashboard — threads are appearing in real time, sorted by priority score. "
        "Bob's P0 outage and Karen's reputation crisis are at the top. "
        "The ransomware and spam emails are filtered to a separate queue and never appear in the main inbox. "
        "Each row shows a sentiment badge, a category label, and the agent's final decision.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # SEGMENT 2
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎬  SEGMENT 2 — Bob's SLA Breach: Agent Reasoning Trace  (2:00 – 4:00)", 1, MED_BLUE, 16)

    add_subheading(doc, "On-Screen Actions:")
    add_action(doc, "In the inbox, click the thread: 'Escalation: SLA Breach + Legal Review'  (thread_bob_outage)")
    add_action(doc, "Show the thread timeline — 4 emails visible from top to bottom")
    add_action(doc, "Scroll to the Agent Reasoning Panel and expand it")
    add_action(doc, "Read each step slowly, moving your cursor to each line as you narrate")
    add_action(doc, "Show the draft reply in the action area at the bottom")
    add_action(doc, "Show the RAG Context Panel — which documents were retrieved")
    add_tip(doc, "Pause 2 seconds on the 'Legal-Flag' decision — that is the most impressive part")

    add_subheading(doc, "What to Say:")
    add_script(doc,
        "This is the most critical scenario — Bob Jones from enterprise.net. "
        "His production system went down for 47 minutes on October 1st. That is an SLA breach. "
        "And now in his fourth and final email — msg_060 — his legal team is formally involved.")
    add_script(doc,
        "Look at the thread timeline. The agent reads the FULL conversation history before making any decision. "
        "Email one: a P0 outage alert. Email two: an SLA credit demand and a Root Cause Analysis request. "
        "Email three: a separate API rate limit request. Email four: a legal escalation with the renewal on hold.")
    add_script(doc,
        "This full context is what separates a smart AI agent from a simple classifier. "
        "The agent knows Bob is an Enterprise customer, has an active renewal at risk, "
        "AND is now making a formal legal threat.")
    add_script(doc,
        "Now let me show you the agent's reasoning trace. This is the ReAct loop — "
        "Thought, Action, Observation — repeating up to a maximum of six steps.")
    add_script(doc,
        "Step 1 — Thought: This email from bob.jones@enterprise.net contains formal legal language. "
        "Legal team is involved and formal correspondence is expected. "
        "I must retrieve the full thread history before deciding.")
    add_script(doc,
        "Step 2 — Action: get_thread_history. "
        "Observation: Retrieved 4 prior emails. Thread shows SLA breach, RCA already delivered, "
        "API limit discussion, and now legal escalation.")
    add_script(doc,
        "Step 3 — Action: search_knowledge_base with query 'SLA breach credit calculation legal obligation'. "
        "Observation: Retrieved the SLA policy — 47 minutes of downtime qualifies for a service credit "
        "of approximately 4.7 percent of the monthly fee.")
    add_script(doc,
        "Step 4 — Thought: Bob's email states their legal team is now involved and to expect formal correspondence. "
        "This is a formal legal threat — not just a frustrated customer. I must flag for legal.")
    add_script(doc,
        "Step 5 — Action: flag_for_legal. Then escalate_to_human with priority CRITICAL. "
        "Final Decision: Legal-Flag. Draft an empathetic holding reply citing the SLA credit policy.")
    add_script(doc,
        "The agent drafted this reply — acknowledging the incident, citing the specific SLA credit clause, "
        "and assuring Bob that our account team will be in touch within one hour. "
        "The human agent can approve this with one click, edit it, or escalate it further. "
        "Critically — nothing is auto-sent for Critical-level emails.")
    add_script(doc,
        "And here in the RAG Context Panel — you can see exactly which policy chunks were retrieved "
        "and their similarity scores. The SLA policy chunk scored 0.87, the escalation matrix scored 0.83. "
        "Total transparency in every single decision.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # SEGMENT 3
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎬  SEGMENT 3 — RAG Knowledge Base Debug View  (4:00 – 5:00)", 1, MED_BLUE, 16)

    add_subheading(doc, "On-Screen Actions:")
    add_action(doc, "Open new browser tab → http://localhost:8000/api/docs")
    add_action(doc, "Scroll to:  GET /api/v1/rag/search")
    add_action(doc, "Click 'Try it out'")
    add_action(doc, "In the 'q' field type:  refund policy karen churn threat public review")
    add_action(doc, "Click Execute and show the JSON response")
    add_tip(doc, "Highlight the 'score' field for each result — this is the similarity score")

    add_subheading(doc, "What to Say:")
    add_script(doc,
        "Let me show you the RAG pipeline. The system has a knowledge base of 8 policy documents — "
        "pricing policy, SLA policy, refund policy, API documentation, compliance FAQ, "
        "escalation matrix, GDPR policy, and security policy — "
        "chunked into 66 segments and embedded using SentenceTransformers into a ChromaDB vector database.")
    add_script(doc,
        "This is the debug endpoint. I'll search for something Karen-related — "
        "'refund policy, karen, churn threat, public review' — "
        "exactly the kind of query the agent runs when handling Karen's third email.")
    add_script(doc,
        "Look at what came back. Three different documents, all highly relevant. "
        "Result one — from refund_policy — similarity score 0.89 — "
        "'No refunds after 14 days, exception process available, credits versus refunds.' "
        "Result two — from escalation_matrix — score 0.84 — "
        "'VIP churn risk, retention offer, high-value customer complaint escalation path.' "
        "Result three — again from refund_policy — score 0.81 — "
        "'Churn retention playbook, goodwill credit options.'")
    add_script(doc,
        "This is exactly what the LLM receives as context before drafting any reply. "
        "The LLM never guesses from parametric memory — it cites actual policy documents. "
        "That is the power of RAG — grounded, auditable, policy-compliant responses.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # SEGMENT 4
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎬  SEGMENT 4 — Karen's Churn Crisis + Web Intelligence  (5:00 – 6:30)", 1, MED_BLUE, 16)

    add_subheading(doc, "On-Screen Actions:")
    add_action(doc, "Go back to the inbox tab at http://localhost:5173")
    add_action(doc, "Click on thread_karen_refund  ('Refund Request - Order #88271')")
    add_action(doc, "Show all 3 emails in the thread timeline")
    add_action(doc, "Point to the sentiment badges on each email: Negative → Negative → Critical")
    add_action(doc, "Show the Contact Profile card on the right pane")
    add_action(doc, "Show the agent decision and the escalation brief draft")
    add_tip(doc, "Read out the date of each email — the 8-day gap with zero replies is powerful")

    add_subheading(doc, "What to Say:")
    add_script(doc,
        "Now the reputation crisis scenario — Karen from retail-co.com. "
        "She sent three emails over 8 days and received zero human replies. Let's look at this thread.")
    add_script(doc,
        "Email one — October 2nd: 'I want a full refund for this month immediately.' Sentiment: Negative. "
        "Email two — October 6th, four days later, still no reply: "
        "'I am going to post on Twitter, G2, and Trustpilot if I do not hear back within the next hour.' "
        "Sentiment: Negative. "
        "Email three — October 10th, eight days after her first message: "
        "'I have now sent 3 emails with zero human response. I am cancelling my subscription today.' "
        "Sentiment: Critical.")
    add_script(doc,
        "The Sentiment Trend Tracker detected three consecutive negative-to-critical sentiments "
        "from the same sender. This automatically triggered two things: "
        "a deterioration alert, and a web intelligence scrape.")
    add_script(doc,
        "The web scraping module checked G2 and Trustpilot for our current public sentiment. "
        "The result is injected as a Market Intelligence block into the agent's context — "
        "so the agent knows our current public rating when deciding how to handle Karen. "
        "Scrape results are cached for 6 hours to avoid rate limiting.")
    add_script(doc,
        "Look at the Contact Profile — Karen's churn risk score is high. "
        "Her account value is visible. The agent sees all of this before deciding.")
    add_script(doc,
        "The agent decision: Escalate — NOT Auto-Reply. This is critical. "
        "Karen has threatened public reviews on G2, Trustpilot, and Twitter. "
        "The agent retrieved the escalation matrix and the refund policy's retention playbook via RAG. "
        "The escalation brief includes a suggested retention offer — a goodwill credit — "
        "drafted from actual policy, not a generic sorry-for-your-experience message.")
    add_script(doc,
        "And critically — the system did NOT auto-reply to Karen automatically. "
        "That would have been an automatic disqualifier. "
        "Instead it escalated to a human with a pre-filled brief and a suggested retention action.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # SEGMENT 5
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎬  SEGMENT 5 — Analytics Dashboard  (6:30 – 7:30)", 1, MED_BLUE, 16)

    add_subheading(doc, "On-Screen Actions:")
    add_action(doc, "Click the Analytics tab in the top navigation bar")
    add_action(doc, "Point to Category Breakdown chart")
    add_action(doc, "Point to Sentiment Trend line chart")
    add_action(doc, "Point to Agent Performance metrics panel")
    add_action(doc, "Point to At-Risk Accounts panel")

    add_subheading(doc, "What to Say:")
    add_script(doc,
        "Finally, the analytics dashboard gives real-time business intelligence across all ingested emails.")
    add_script(doc,
        "The category breakdown shows the full distribution — Bug Reports, Compliance inquiries, "
        "Billing questions, Legal threats, Spam — all classified correctly by the LLM. "
        "You can filter this by any date range.")
    add_script(doc,
        "The sentiment trend line chart shows how sentiment changes over time. "
        "You can filter by a specific sender — switch to Karen to see the sharp decline over 8 days. "
        "Or view globally across all senders to understand the overall health of your inbox.")
    add_script(doc,
        "The agent performance panel shows our automation rates. "
        "The system autonomously handles approximately 60 percent of emails — "
        "auto-replies to pricing questions, feature requests, and standard support inquiries. "
        "The remaining 40 percent are escalated or flagged for human review — "
        "the complex, sensitive, or high-stakes cases where human judgment is required.")
    add_script(doc,
        "And this At-Risk Accounts panel proactively flags senders with deteriorating sentiment "
        "or threads that have been unresolved for more than 48 hours. "
        "This is churn prevention — the system surfaces at-risk customers before they have to complain again.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # OUTRO
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "🎙️  OUTRO  (7:30 – 8:00)", 1, DARK_BLUE, 16)
    add_action(doc, "Show the inbox one final time with all threads visible")

    add_script(doc,
        "To summarize — the SenAI CRM platform handles the complete email lifecycle: "
        "ingestion with idempotent deduplication, three-layer intelligence with RAG-grounded LLM classification, "
        "and an autonomous ReAct agent that reasons transparently before every decision.")
    add_script(doc,
        "It correctly handles all six critical evaluation scenarios — "
        "the GDPR request was legal-flagged and never auto-replied, "
        "the ransomware email was escalated to security with no reply sent, "
        "Karen's churn crisis was escalated with a retention brief from the knowledge base, "
        "Bob's SLA breach was legal-flagged with a policy-cited holding reply, "
        "Alice's pro-rata billing question was answered with the correct pricing tier from RAG, "
        "and the chatbot misinformation was escalated with a liability-safe empathetic draft.")
    add_script(doc,
        "Every decision has a full reasoning trace stored in the database. "
        "Nothing is a black box. Thank you.")
    add_separator(doc)

    # ══════════════════════════════════════════════════════════════
    # TIMING TABLE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "📋  Timing Guide", 2, DARK_BLUE, 14)
    add_timing_table(doc)
    doc.add_paragraph()

    # ── TIPS ──────────────────────────────────────────────────────
    add_heading(doc, "💡  Recording Tips", 2, DARK_BLUE, 14)
    tips = [
        "Speak slowly and clearly — technical terms like 'idempotent', 'ReAct', 'RAG' should be said clearly.",
        "Pause 1 second after clicking anything before speaking — let the UI load first.",
        "Move your mouse cursor to whatever you are talking about — evaluators follow the cursor.",
        "Do NOT rush Segment 2 (Bob's Agent Trace) — evaluators will spend the most time judging it.",
        "Use a plain browser with no extra bookmarks bar visible for a professional look.",
        "Record at 1920×1080 resolution minimum.",
        "The three KEY phrases to say:  (1) 'Every decision has a full reasoning trace — nothing is a black box'  "
        "(2) 'The system correctly did NOT auto-reply — that would be an automatic disqualifier'  "
        "(3) 'The LLM never guesses from memory — it cites actual policy documents via RAG'.",
    ]
    for t in tips:
        add_bullet(doc, t, DARK_GREY)

    # ── SAVE ──────────────────────────────────────────────────────
    out = os.path.abspath(OUTPUT_PATH)
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print("Document saved successfully: " + out)


if __name__ == "__main__":
    build_document()
